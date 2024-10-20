from flask import Flask, render_template, request, redirect, url_for, session, send_file
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
from docx import Document

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Секретный ключ для работы сессией
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///words.db'  # Путь к базе данных
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Отключаем уведомления о изменениях
db = SQLAlchemy(app)

# Модель пользователя
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    words = db.relationship('Word', backref='user', lazy=True)

# Модель слов
class Word(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.String(50), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

# Создание базы данных
with app.app_context():
    db.create_all()  # Создает все таблицы

# Главная страница
@app.route('/')
def home():
    if 'user_id' in session:
        return redirect(url_for('add_words'))
    return redirect(url_for('login'))

# Страница регистрации
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if User.query.filter_by(username=username).first():
            return 'Пользователь уже существует. Попробуйте другой логин.'
        new_user = User(username=username, password=password)
        db.session.add(new_user)
        db.session.commit()
        return redirect(url_for('login'))
    return render_template('register.html')

# Страница входа
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            session['user_id'] = user.id  # Сохраняем ID пользователя в сессии
            return redirect(url_for('add_words'))
        else:
            return 'Неверные учетные данные!'
    return render_template('login.html')

# Страница для добавления слов
@app.route('/add_words', methods=['GET', 'POST'])
def add_words():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        words = request.form['words']
        new_words = [word.strip() for word in words.replace(',', ' ').split()]
        user_id = session['user_id']
        for word in new_words:
            if word:  # Проверка на пустые строки
                new_word = Word(content=word, user_id=user_id)
                db.session.add(new_word)
        db.session.commit()
        return redirect(url_for('filter_words'))
    return render_template('add_words.html')

# Страница с отфильтрованными иероглифами
@app.route('/filter_words')
def filter_words():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    user_words = Word.query.filter_by(user_id=user_id).order_by(Word.id.desc()).all()  # Получаем слова, сортируя по ID в обратном порядке

    # Уникальные иероглифы в порядке добавления (новые сверху)
    seen = set()
    unique_characters = []
    for word in user_words:
        for char in word.content:
            if char not in seen:
                seen.add(char)
                unique_characters.append(char)

    return render_template('filtered.html', characters=unique_characters, enumerate=enumerate)

# Очистка слов пользователя
@app.route('/clear')
def clear():
    if 'user_id' in session:
        user_id = session['user_id']
        Word.query.filter_by(user_id=user_id).delete()  # Удаляем все слова пользователя
        db.session.commit()
    return redirect(url_for('add_words'))

# Выход из системы
@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('login'))

# Функция для скачивания уникальных иероглифов в Word
@app.route('/download')
def download():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    user_words = Word.query.filter_by(user_id=user_id).order_by(Word.id.desc()).all()  # Получаем слова, сортируя по ID в обратном порядке

    # Создаем документ Word
    doc = Document()
    doc.add_heading('Уникальные иероглифы', level=1)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'Иероглифы'

    seen = set()
    unique_characters = []
    for word in user_words:
        for char in word.content:
            if char not in seen:
                seen.add(char)
                unique_characters.append(char)

    # Заполняем таблицу уникальными иероглифами
    for char in unique_characters:
        row_cells = table.add_row().cells
        row_cells[0].text = char

    # Сохраняем документ в байтовом потоке
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    # Возвращаем файл для скачивания
    return send_file(bio, as_attachment=True, download_name='unique_characters.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


